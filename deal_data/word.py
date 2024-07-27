import pandas as pd
from docx import Document


def csv_to_programming_questions(csv_file, word_file):
    # 读取 CSV 文件为 DataFrame
    df = pd.read_csv(csv_file)

    # 创建一个新的 Word 文档
    doc = Document()

    # 遍历 DataFrame 的每一行数据
    for idx, row in df.iterrows():
        # 构建编程题格式的文本
        question_text = f"{idx}. {row['problem_text']}\n"
        question_text += f"知识点: {row['skill_name']} \n"
        question_text += f"难度等级: {row['level']} \n"
        if pd.notna(row['answer']):
            question_text += f"答案: {row['answer']} \n"

        # 添加问题到 Word 文档中
        doc.add_paragraph(question_text)
        doc.add_paragraph("")  # 添加空行分隔每个问题

    # 保存 Word 文档
    doc.save(word_file)

    print(f'CSV 文件 {csv_file} 已成功转换为 Word 文件 {word_file}')


# 示例用法
csv_file = './program_question.csv'  # 输入你的 CSV 文件名
word_file = 'programming.docx'  # 输出 Word 文件名

csv_to_programming_questions(csv_file, word_file)
